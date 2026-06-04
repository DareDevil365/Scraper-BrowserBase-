#!/usr/bin/env python3
"""
present.py â€” Deterministic presentation/rendering layer (Stage 3, no API calls).

Referenced by Synthesis.md Â§8 (Presentation Layer). Standalone: run by hand with
    python present.py /path/to/run_<id> [--formats docx,html,pdf] [--no-charts]

CONTRACTS (kept identical to Synthesis.md Â§1 and fallback_synth.py):
  * Reads ONLY from disk. Never re-scrapes, never calls a model.
  * Input precedence for the body it presents:
        final_report_v2.txt  >  final_report_v1.txt  >  final_report_fallback.txt
    Whichever exists (richest first) is the prose body.
  * Structured visuals (tables / charts / timeline) are promoted from
    extracted/<vector_id>.json â€” the same payloads synthesis used.
  * Coverage dashboard + sources table come from state.json + sources.json.
  * If the body was the fallback file, the "MECHANICALLY ASSEMBLED" banner is
    preserved and re-stamped at the top of every output format.

OUTPUT (multi-format, one render pass):
    report.html   (always; self-contained, charts inlined as base64)
    report.docx   (via python-docx)
    report.pdf    (via reportlab; pandoc fallback)

Nothing here overwrites v1/v2/fallback markdown â€” it only adds report.* files.
"""

from __future__ import annotations

import argparse
import base64
import datetime as _dt
import html as _html
import io
import json
import os
import re
import sys
from pathlib import Path
from typing import Any

# ----- optional deps, all degrade gracefully ---------------------------------
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _HAVE_MPL = True
except Exception:
    _HAVE_MPL = False

try:
    import markdown as _md
    _HAVE_MD = True
except Exception:
    _HAVE_MD = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAVE_DOCX = True
except Exception:
    _HAVE_DOCX = False

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors as _rl_colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    )
    _HAVE_REPORTLAB = True
except Exception:
    _HAVE_REPORTLAB = False

import shutil
_HAVE_PANDOC = shutil.which("pandoc") is not None

try:
    import jinja2 as _jinja2
    _HAVE_JINJA = True
except Exception:
    _HAVE_JINJA = False

DEFAULT_CONFIG: dict[str, Any] = {
    "formats": {
        "default": [
            "docx",
            "html",
            "pdf"
        ],
        "pdf_engine_order": [
            "reportlab",
            "pandoc"
        ]
    },
    "body_precedence": {
        "order": [
            "final_report_v2.txt",
            "final_report_v1.txt",
            "final_report_fallback.txt",
            "partial_report.txt"
        ]
    },
    "visual_detection": {
        "timeline": {
            "min_events": 2,
            "date_keys": [
                "date",
                "year",
                "when",
                "time",
                "timestamp"
            ],
            "label_keys": [
                "event",
                "description",
                "milestone",
                "title"
            ]
        },
        "table": {
            "min_rows": 2,
            "min_shared_keys": 2,
            "max_shared_keys": 8
        },
        "bar_chart": {
            "min_numeric_values": 3
        }
    },
    "coverage": {
        "prefer_state_json": True,
        "low_coverage_max_leaves": 2,
        "badges": {
            "RICH": "🟢 RICH",
            "PARTIAL": "🟡 PARTIAL",
            "LOW_COVERAGE": "⚠️ LOW COVERAGE",
            "EMPTY": "⚪ EMPTY"
        }
    },
    "sources": {
        "dedupe_by": "url",
        "group_by": "vector_id",
        "junk_marker": "⚠ off-topic",
        "junk_domain_substrings": [
            "dictionary.cambridge.org",
            "merriam-webster.com",
            "thefreedictionary.com",
            "britannica.com/dictionary",
            "apdpms.ap.gov.in",
            "wiktionary.org",
            "translate.google"
        ]
    },
    "banner": {
        "fallback_text": "⚠️ MECHANICALLY ASSEMBLED — no AI synthesis ran; structured from extracted data only.",
        "trigger_synthesis_modes": [
            "fallback_no_api",
            "incomplete_fallback"
        ]
    }
}


def _merge_dicts(default: dict, override: dict) -> dict:
    res = dict(default)
    for k, v in override.items():
        if isinstance(v, dict) and k in res and isinstance(res[k], dict):
            res[k] = _merge_dicts(res[k], v)
        else:
            res[k] = v
    return res


def load_present_config() -> dict[str, Any]:
    base = Path(__file__).parent
    for name in ["present_config.json", "Present_config.json"]:
        config_path = base / name
        if config_path.exists():
            try:
                override = json.loads(config_path.read_text(encoding="utf-8"))
                return _merge_dicts(DEFAULT_CONFIG, override)
            except Exception as e:
                print(f"Warning: Failed to load {name}: {e}", file=sys.stderr)
    return DEFAULT_CONFIG


CONFIG = load_present_config()

# =============================================================================
# 1. Disk loading (read-only)
# =============================================================================
BODY_PRECEDENCE = CONFIG.get("body_precedence", {}).get("order", [
    "final_report_v2.txt",
    "final_report_v1.txt",
    "final_report_fallback.txt",
    "partial_report.txt",
])

FALLBACK_BANNER_RE = re.compile(r"mechanically assembled", re.I)


def load_json_if_exists(path: Path, default: Any = None) -> Any:
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            pass
    return default


def load_best_payload(run_dir: Path) -> dict[str, Any]:
    # Determine best markdown body
    body_path, body_text = None, ""
    for name in BODY_PRECEDENCE:
        p = run_dir / name
        if p.exists():
            body_path, body_text = p, p.read_text(encoding="utf-8", errors="replace")
            break

    is_fallback = bool(body_path and "fallback" in body_path.name) or \
        bool(FALLBACK_BANNER_RE.search(body_text[:2000]))

    # Parse markdown into a structured list of sections
    parsed_sections = []
    title = ""
    executive_summary = ""

    if body_text and not is_fallback:
        lines = body_text.splitlines()
        current_section_title = ""
        current_section_body = []
        
        for line in lines:
            m_title = re.match(r"^#\s+(.*)", line)
            if m_title and not title:
                title = m_title.group(1).strip()
                continue
                
            m_sec = re.match(r"^##\s+(.*)", line)
            if m_sec:
                if current_section_title:
                    sec_body_str = "\n".join(current_section_body).strip()
                    if current_section_title.lower() in ("executive summary", "summary"):
                        executive_summary = sec_body_str
                    else:
                        parsed_sections.append({
                            "title": current_section_title,
                            "body": sec_body_str
                        })
                current_section_title = m_sec.group(1).strip()
                current_section_body = []
                continue
                
            current_section_body.append(line)
            
        if current_section_title:
            sec_body_str = "\n".join(current_section_body).strip()
            if current_section_title.lower() in ("executive summary", "summary"):
                executive_summary = sec_body_str
            else:
                parsed_sections.append({
                    "title": current_section_title,
                    "body": sec_body_str
                })

    elif body_text and is_fallback:
        m_title = re.match(r"^#\s+(.*)", body_text.splitlines()[0] if body_text.splitlines() else "")
        if m_title:
            title = m_title.group(1).strip()

    state_json = load_json_if_exists(run_dir / "state.json", {})
    run_config = load_json_if_exists(run_dir / "run_config.json", {})
    
    # Get all vectors from state.json
    vectors = state_json.get("vectors") or []
    
    # We will build raw_sections combining parsed_sections and vectors
    raw_sections = []
    matched_vector_ids = set()
    
    # Helper to clean titles for comparison
    def clean_t(t):
        return re.sub(r"[^\w]+", "", str(t).lower())

    # Try matching parsed_sections to state.json vectors
    for p_sec in parsed_sections:
        sec_title_clean = clean_t(p_sec["title"])
        matched_vec = None
        
        for vec in vectors:
            vec_id = vec.get("id") or vec.get("vector_id")
            if vec_id in matched_vector_ids:
                continue
            vec_topic = vec.get("topic") or vec.get("title") or ""
            if clean_t(vec_topic) == sec_title_clean or sec_title_clean in clean_t(vec_topic) or clean_t(vec_topic) in sec_title_clean:
                matched_vec = vec
                break
                
        if matched_vec:
            vec_id = matched_vec.get("id") or matched_vec.get("vector_id")
            matched_vector_ids.add(vec_id)
            # Load extracted data
            data_file = run_dir / "extracted" / f"{vec_id}.json"
            extracted_data = load_json_if_exists(data_file)
            
            raw_sections.append({
                "vector_id": vec_id,
                "topic": matched_vec.get("topic") or p_sec["title"],
                "body": p_sec["body"],
                "data": extracted_data
            })
        else:
            raw_sections.append({
                "vector_id": "",
                "topic": p_sec["title"],
                "body": p_sec["body"],
                "data": None
            })
            
    # Now add any vectors that weren't matched
    for vec in vectors:
        vec_id = vec.get("id") or vec.get("vector_id")
        if vec_id not in matched_vector_ids:
            data_file = run_dir / "extracted" / f"{vec_id}.json"
            extracted_data = load_json_if_exists(data_file)
            raw_sections.append({
                "vector_id": vec_id,
                "topic": vec.get("topic") or vec_id,
                "body": "",
                "data": extracted_data
            })
            
    if not title:
        title = run_config.get("query") or "Research Report"
        
    payload = {
        "title": title,
        "session_id": run_config.get("session_id") or state_json.get("session_id") or run_dir.name,
        "created_at": state_json.get("created_at") or run_config.get("created_at") or "",
        "status": state_json.get("status") or "complete",
        "executive_summary": executive_summary,
        "sections": raw_sections,
        "is_fallback": is_fallback,
        "synthesis_mode": state_json.get("synthesis_mode", "ai"),
        "original_query": run_config.get("query") or state_json.get("query") or "",
        "clarification_answers": run_config.get("clarification_answers") or state_json.get("clarification_answers") or [],
    }
    return payload


def _read_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


# =============================================================================
# 2. Structured-data detection  (promote JSON -> tables / charts / timeline)
# =============================================================================
NUM_RE = re.compile(r"(-?\d[\d,]*\.?\d*)")


def _to_number(v: Any) -> float | None:
    if isinstance(v, (int, float)):
        return float(v)
    if isinstance(v, str):
        m = NUM_RE.search(v.replace(",", ""))
        if m:
            try:
                return float(m.group(1))
            except ValueError:
                return None
    return None


def detect_visuals(payload: Any) -> list[dict[str, Any]]:
    """
    Walk a payload and return a list of visual specs:
      {"kind": "timeline", "title", "events":[{date,label}]}
      {"kind": "table",    "title", "headers", "rows"}
      {"kind": "barchart", "title", "labels", "values"}
    Pure heuristic; if nothing qualifies, returns []. No inference.
    """
    visuals: list[dict[str, Any]] = []
    if not isinstance(payload, (dict, list)):
        return visuals

    vis_cfg = CONFIG.get("visual_detection", {})
    tl_cfg = vis_cfg.get("timeline", {})
    tbl_cfg = vis_cfg.get("table", {})
    bc_cfg = vis_cfg.get("bar_chart", {})

    min_events = tl_cfg.get("min_events", 2)
    date_keys = tuple(tl_cfg.get("date_keys", ["date", "year", "when", "time", "timestamp"]))
    label_keys = tuple(tl_cfg.get("label_keys", ["event", "description", "milestone", "title", "label", "name"]))

    min_rows = tbl_cfg.get("min_rows", 2)
    min_shared_keys = tbl_cfg.get("min_shared_keys", 2)
    max_shared_keys = tbl_cfg.get("max_shared_keys", 8)

    min_numeric_values = bc_cfg.get("min_numeric_values", 3)

    def walk(node: Any, keyname: str = ""):
        # Timeline: list of dicts each carrying a date/year + an event/label
        if isinstance(node, list) and node and all(isinstance(x, dict) for x in node):
            sample = node[0]
            date_key = _first_key(sample, date_keys)
            label_key = _first_key(sample, label_keys)
            if date_key and label_key:
                events = []
                for x in node:
                    d, l = x.get(date_key), x.get(label_key)
                    if d and l:
                        events.append({"date": str(d), "label": str(l)})
                if len(events) >= min_events:
                    visuals.append({"kind": "timeline",
                                    "title": _pretty(keyname) or "Timeline",
                                    "events": events})
            # Object table: list of flat dicts with shared scalar keys
            flat = [x for x in node if _is_flat(x)]
            if len(flat) >= min_rows:
                headers = _union_keys(flat)
                if min_shared_keys <= len(headers) <= max_shared_keys:
                    rows = [[_cell(x.get(h, "")) for h in headers] for x in flat]
                    visuals.append({"kind": "table",
                                    "title": _pretty(keyname) or "Details",
                                    "headers": [_pretty(h) for h in headers],
                                    "rows": rows})
            for i, x in enumerate(node):
                walk(x, f"{keyname}")
            return

        if isinstance(node, dict):
            # Bar chart: a flat dict whose values are mostly numeric (>=min_numeric_values numeric)
            numeric = {k: _to_number(v) for k, v in node.items()
                       if not isinstance(v, (dict, list))}
            numeric = {k: v for k, v in numeric.items() if v is not None}
            if len(numeric) >= min_numeric_values:
                visuals.append({"kind": "barchart",
                                "title": _pretty(keyname) or "Figures",
                                "labels": [_pretty(k) for k in numeric],
                                "values": list(numeric.values())})
            for k, v in node.items():
                walk(v, k)

    walk(payload, "")
    # De-dupe by (kind,title) keeping the first/richest
    seen, out = set(), []
    for v in visuals:
        sig = (v["kind"], v["title"])
        if sig not in seen:
            seen.add(sig)
            out.append(v)
    return out


def _first_key(d: dict, candidates: tuple[str, ...]) -> str | None:
    lower = {k.lower(): k for k in d}
    for c in candidates:
        if c in lower:
            return lower[c]
    return None


def _is_flat(d: Any) -> bool:
    return isinstance(d, dict) and all(
        not isinstance(v, (dict, list)) for v in d.values()
    )


def _union_keys(dicts: list[dict]) -> list[str]:
    keys: list[str] = []
    for d in dicts:
        for k in d:
            if k not in keys:
                keys.append(k)
    return keys


def _cell(v: Any) -> str:
    s = str(v)
    return (s[:140] + "â€¦") if len(s) > 140 else s


def _pretty(s: str) -> str:
    if not s:
        return ""
    s = re.sub(r"[_\-]+", " ", str(s))
    return s.strip().title()


# =============================================================================
# 3. Coverage dashboard + sources table
# =============================================================================
RICH, PARTIAL, LOW, EMPTY = "RICH", "PARTIAL", "LOW_COVERAGE", "EMPTY"


def grade_payload(payload: Any) -> str:
    if payload is None:
        return EMPTY
    leaves = _count_leaves(payload)
    if leaves == 0:
        return EMPTY
    cov_cfg = CONFIG.get("coverage", {})
    low_max = cov_cfg.get("low_coverage_max_leaves", 2)
    if leaves <= low_max:
        return LOW
    if leaves <= 5:
        return PARTIAL
    return RICH


def _count_leaves(node: Any) -> int:
    if isinstance(node, dict):
        return sum(_count_leaves(v) for v in node.values())
    if isinstance(node, list):
        return sum(_count_leaves(v) for v in node)
    if node in (None, "", "null"):
        return 0
    return 1


def build_coverage(data: dict[str, Any]) -> dict[str, Any]:
    state = data.get("state", {})
    extracted = data.get("extracted", {})
    cov_cfg = CONFIG.get("coverage", {})
    prefer_state = cov_cfg.get("prefer_state_json", True)
    
    # Prefer explicit per-vector status from state.json if present.
    per_vector = state.get("vector_status") or {} if prefer_state else {}
    
    # Build a lookup mapping vector_id -> topic from state.json
    state_vectors = {}
    for v in state.get("vectors", []):
        if isinstance(v, dict) and "id" in v:
            state_vectors[v["id"]] = v.get("topic") or v["id"]
            
    rows = []
    badges_map = cov_cfg.get("badges", {})
    for vid, payload in sorted(extracted.items()):
        status = per_vector.get(vid) or grade_payload(payload)
        topic = state_vectors.get(vid, vid)
        badge = badges_map.get(status, status)
        rows.append({
            "vector_id": vid,
            "topic": topic,
            "status": status,
            "badge": badge
        })
    total = state.get("total_vectors") or len(rows) or 0
    completed = state.get("completed_vectors")
    if completed is None:
        completed = sum(1 for r in rows if r["status"] in (RICH, PARTIAL))
    return {
        "rows": rows,
        "completed": completed,
        "total": total,
        "synthesis_mode": state.get("synthesis_mode", "ai"),
        "status": state.get("status", "complete"),
        "counts": {
            RICH: sum(1 for r in rows if r["status"] == RICH),
            PARTIAL: sum(1 for r in rows if r["status"] == PARTIAL),
            LOW: sum(1 for r in rows if r["status"] == LOW),
            EMPTY: sum(1 for r in rows if r["status"] == EMPTY),
        },
    }


JUNK_HINTS = ("dictionary", "cambridge.org", "merriam-webster", "thefreedictionary",
              "britannica.com/dictionary", "apdpms.ap.gov.in", "building permit")


def build_sources(data: dict[str, Any]) -> list[dict[str, Any]]:
    seen, out = set(), []
    src_cfg = CONFIG.get("sources", {})
    junk_hints = src_cfg.get("junk_domain_substrings", JUNK_HINTS)
    dedupe_key = src_cfg.get("dedupe_by", "url")
    group_key = src_cfg.get("group_by", "vector_id")
    
    for s in data.get("sources", []):
        url = (s.get("url") or "").strip()
        if not url:
            continue
        dup_val = (s.get(dedupe_key) or url).strip() if dedupe_key else url
        if not dup_val or dup_val in seen:
            continue
        seen.add(dup_val)
        title = (s.get("title") or "").strip()
        looks_junk = any(h in url.lower() or h in title.lower() for h in junk_hints)
        out.append({
            "url": url,
            "title": title or url,
            "tier": s.get("tier", ""),
            "status": s.get("status", ""),
            "vector_id": s.get("vector_id", ""),
            "flagged": looks_junk,
            "junk": looks_junk
        })
    # group by vector_id/group_key, keep flagged at bottom
    out.sort(key=lambda r: (r["flagged"], r.get(group_key, "") if group_key else "", r["title"].lower()))
    return out


# =============================================================================
# 4. Chart rendering (matplotlib -> PNG bytes); degrade to None if unavailable
# =============================================================================
def render_barchart(spec: dict[str, Any]) -> bytes | None:
    if not _HAVE_MPL:
        return None
    labels, values = spec["labels"], spec["values"]
    pairs = sorted(zip(labels, values), key=lambda p: p[1], reverse=True)[:12]
    labels, values = [p[0] for p in pairs], [p[1] for p in pairs]
    fig, ax = plt.subplots(figsize=(7.2, max(2.4, 0.45 * len(labels))))
    ax.barh(range(len(labels)), values, color="#3b6ea5")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=8)
    ax.invert_yaxis()
    ax.set_title(spec["title"], fontsize=10)
    for i, v in enumerate(values):
        ax.text(v, i, f" {v:g}", va="center", fontsize=7)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    return buf.getvalue()


# =============================================================================
# 5. Assemble an intermediate "document model" then render each format
# =============================================================================
def build_document_model(data: dict[str, Any]) -> dict[str, Any]:
    cfg = data.get("config", {})
    cov = build_coverage(data)
    sources = build_sources(data)

    # gather visuals across all extracted payloads
    visuals: list[dict[str, Any]] = []
    for vid, payload in sorted(data["extracted"].items()):
        for v in detect_visuals(payload):
            v = dict(v)
            v["vector_id"] = vid
            visuals.append(v)

    title = (cfg.get("query") or "Research Report").strip()
    title = (title[:120] + "â€¦") if len(title) > 120 else title

    return {
        "title": title,
        "generated_at": _dt.datetime.now().isoformat(timespec="seconds"),
        "session_id": cfg.get("session_id", data["run_dir"].name),
        "is_fallback": data["is_fallback"],
        "coverage": cov,
        "sources": sources,
        "visuals": visuals,
        "body_md": data["body_text"],
        "body_origin": data["body_path"].name if data["body_path"] else None,
    }


def _banner_text(model: dict[str, Any]) -> str | None:
    cov = model["coverage"]
    banner_cfg = CONFIG.get("banner", {})
    fallback_txt = banner_cfg.get("fallback_text", "⚠️ MECHANICALLY ASSEMBLED — no AI synthesis ran; structured from extracted data only.")
    trigger_modes = banner_cfg.get("trigger_synthesis_modes", ["fallback_no_api", "incomplete_fallback"])
    
    if model["is_fallback"] or cov["synthesis_mode"] in trigger_modes:
        try:
            return fallback_txt.format(completed=cov['completed'], total=cov['total'])
        except Exception:
            return fallback_txt
    if cov["status"] not in ("complete", ""):
        return (f"⚠️ INCOMPLETE — {cov['completed']}/{cov['total']} vectors "
                f"(status: {cov['status']}).")
    return None


# ---- 5a. HTML -------------------------------------------------------------
# ---- 5a. HTML -------------------------------------------------------------
def assert_no_template_or_placeholder_leak(content: str):
    # Check for unrendered Jinja tags
    if "{{" in content or "}}" in content or "{%" in content or "%}" in content:
        raise ValueError("Quality Gate Violation: Unrendered Jinja template tags found in final report.")
        
    # Check for forbidden placeholder phrases
    from presentation_filter import PLACEHOLDER_PATTERNS
    for pattern in PLACEHOLDER_PATTERNS:
        if pattern.lower() in content.lower():
            raise ValueError(f"Quality Gate Violation: Placeholder / error pattern '{pattern}' leaked into final report.")


def _md_to_html(md_text: str) -> str:
    """Fallback basic markdown-to-HTML parser for tables, lists, headers, and bold text."""
    lines = md_text.splitlines()
    html_parts = []
    in_list = False
    in_table = False
    table_headers = None
    table_rows = []

    def flush_list():
        nonlocal in_list
        if in_list:
            html_parts.append("</ul>")
            in_list = False

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if in_table:
            html_parts.append("<table>")
            if table_headers:
                html_parts.append("<thead><tr>" + "".join(f"<th>{_html.escape(h)}</th>" for h in table_headers) + "</tr></thead>")
            html_parts.append("<tbody>")
            for row in table_rows:
                html_parts.append("<tr>" + "".join(f"<td>{_html.escape(cell)}</td>" for cell in row) + "</tr>")
            html_parts.append("</tbody></table>")
            in_table = False
            table_headers = None
            table_rows = []

    for line in lines:
        stripped = line.strip()
        
        # Table row matching
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_list()
            parts = [p.strip() for p in stripped.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_headers = parts
            else:
                # Check if it's separator line (e.g. |---|---|)
                if all(re.match(r"^:?-+:?$", cell) for cell in parts if cell.strip()):
                    continue
                table_rows.append(parts)
            continue
        else:
            flush_table()

        # Headers matching
        m_header = re.match(r"^(#{1,6})\s+(.*)", line)
        if m_header:
            flush_list()
            level = len(m_header.group(1))
            content = m_header.group(2).strip()
            html_parts.append(f"<h{level}>{_html.escape(content)}</h{level}>")
            continue

        # List items matching
        m_list = re.match(r"^\s*[-*]\s+(.*)", line)
        if m_list:
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            content = m_list.group(1).strip()
            content = _html.escape(content)
            content = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", content)
            content = re.sub(r"\*(.*?)\*", r"<em>\1</em>", content)
            html_parts.append(f"<li>{content}</li>")
            continue
        else:
            flush_list()

        # Normal text paragraph or empty line
        if stripped:
            content = _html.escape(stripped)
            content = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", content)
            content = re.sub(r"\*(.*?)\*", r"<em>\1</em>", content)
            html_parts.append(f"<p>{content}</p>")
        else:
            html_parts.append("<br/>")

    flush_list()
    flush_table()
    return "\n".join(html_parts)


def render_html_template(presentable: dict[str, Any]) -> str:
    template_path = Path(__file__).parent / "templates" / "Report_Template.html"
    if not template_path.exists():
        template_path = Path(__file__).parent / "templates" / "report_template.html"
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found at {template_path}")
        
    # Render markdown to HTML for Jinja2 safe display
    import copy
    rendered_presentable = copy.deepcopy(presentable)
    for sec in rendered_presentable.get("sections", []):
        body_md = sec.get("body", "")
        
        # Replace markdown images in body_md first
        body_md = re.sub(
            r'!\[(.*?)\]\((.*?)\)',
            r'<div class="image-wrapper"><img class="report-img" src="\2" alt="\1"/><div class="image-fallback"></div></div>',
            body_md
        )
        
        if _HAVE_MD:
            import markdown as _md
            sec["body"] = _md.markdown(body_md, extensions=["tables", "fenced_code"])
        else:
            sec["body"] = _md_to_html(body_md)
            
        # Post-process body HTML to render Mermaid code blocks properly
        import html as std_html
        def unescape_mermaid(match):
            escaped_code = match.group(1)
            # Unescape HTML entities inside mermaid block
            unescaped_code = std_html.unescape(escaped_code)
            # Strip outer tags if any (like <code>)
            unescaped_code = re.sub(r'^<code[^>]*>|</code>$', '', unescaped_code, flags=re.DOTALL)
            return f'<div class="mermaid">{unescaped_code}</div>'
            
        sec["body"] = re.sub(
            r'<pre><code class="language-mermaid">(.*?)</code></pre>',
            unescape_mermaid,
            sec["body"],
            flags=re.DOTALL
        )
        sec["body"] = re.sub(
            r'<pre class="mermaid">(.*?)</pre>',
            unescape_mermaid,
            sec["body"],
            flags=re.DOTALL
        )
        
        # Also wrap any existing raw HTML img tags that might have slipped through
        sec["body"] = re.sub(
            r'<img\s+([^>]*?)src="([^"]+)"([^>]*?)>',
            r'<div class="image-wrapper"><img class="report-img" src="\2" \1 \3/><div class="image-fallback"></div></div>',
            sec["body"]
        )
            
    # Set up jinja2 environment using FileSystemLoader and select_autoescape
    env = _jinja2.Environment(
        loader=_jinja2.FileSystemLoader(str(template_path.parent)),
        autoescape=_jinja2.select_autoescape(["html", "xml"]),
    )
    template = env.get_template(template_path.name)
    
    html_out = template.render(**rendered_presentable)
    return html_out


# ---- 5b. DOCX -------------------------------------------------------------
# ---- 5b. DOCX -------------------------------------------------------------
def render_docx(presentable: dict[str, Any], out_path: Path) -> bool:
    if not _HAVE_DOCX:
        return False
        
    doc = Document()
    doc.add_heading(presentable["title"], level=0)
    
    meta = doc.add_paragraph()
    meta.add_run(f"Session {presentable['session_id']} · generated "
                 f"{presentable['generated_at']} · status: {presentable['status']}").italic = True

    banner = presentable.get("banner")
    if banner:
        p = doc.add_paragraph()
        run = p.add_run(banner)
        run.bold = True
        run.font.color.rgb = RGBColor(0xB0, 0x6A, 0x00)

    # Executive Summary
    exec_sum = presentable.get("executive_summary")
    if exec_sum:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(exec_sum)

    # Key Takeaways
    takeaways = presentable.get("key_takeaways")
    if takeaways:
        doc.add_heading("Key Takeaways", level=2)
        for item in takeaways:
            doc.add_paragraph(item, style="List Bullet")

    # Sections
    for sec in presentable.get("sections", []):
        doc.add_heading(sec["title"], level=1)
        _md_to_docx(doc, sec.get("body", ""), out_path.parent / "media")
        if sec.get("chart_bytes"):
            try:
                import io
                doc.add_picture(io.BytesIO(sec["chart_bytes"]), width=Inches(6.0))
            except Exception as e:
                print(f"  [docx] Failed to add chart image: {e}", file=sys.stderr)

    # Tables
    tables = presentable.get("tables", [])
    if tables:
        doc.add_heading("Extracted Parameter Matrix", level=1)
        for table in tables:
            doc.add_heading(table["title"], level=2)
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if not rows:
                continue
                
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Light Grid Accent 1"
            for i, h in enumerate(headers):
                t.rows[0].cells[i].text = h
                
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    if i < len(cells):
                        cells[i].text = str(val or "")

    # Coverage Gaps
    gaps = presentable.get("gaps", [])
    if gaps:
        doc.add_heading("Coverage Gaps", level=1)
        for gap in gaps:
            doc.add_paragraph(f"{gap['topic']}: {gap['reason']}", style="List Bullet")

    # Sources
    sources = presentable.get("sources", [])
    if sources:
        doc.add_heading("Sources & References", level=1)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = "Source Title"
        t.rows[0].cells[1].text = "Domain"
        t.rows[0].cells[2].text = "URL"
        
        for s in sources:
            c = t.add_row().cells
            c[0].text = str(s.get("title") or "")
            c[1].text = str(s.get("domain") or "")
            c[2].text = str(s.get("url") or "")

    try:
        doc.save(str(out_path))
        return True
    except PermissionError:
        print(f"Error: Permission denied when writing to {out_path}. Please close the file if it is open in another application.")
        fallback_path = out_path.with_name(f"{out_path.stem}_fallback{out_path.suffix}")
        try:
            doc.save(str(fallback_path))
            print(f"Successfully saved to fallback path: {fallback_path}")
            return True
        except Exception as ex:
            print(f"Failed to write to fallback path: {ex}")
            raise
    except Exception as e:
        print(f"Failed to save DOCX: {e}")
        raise


def download_image_cached(url: str, media_dir: Path) -> Path | None:
    if not url.startswith("http"):
        return None
    try:
        import urllib.request
        import hashlib
        hashed = hashlib.md5(url.encode()).hexdigest()[:12]
        ext = ".jpg"
        if ".png" in url.lower():
            ext = ".png"
        elif ".gif" in url.lower():
            ext = ".gif"
        elif ".webp" in url.lower():
            ext = ".webp"
            
        local_path = media_dir / f"img_{hashed}{ext}"
        if local_path.exists():
            return local_path
            
        media_dir.mkdir(parents=True, exist_ok=True)
        # Download image with user-agent
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        )
        with urllib.request.urlopen(req, timeout=10) as response:
            with open(local_path, "wb") as f:
                f.write(response.read())
        return local_path
    except Exception as e:
        print(f"Failed to download image {url}: {e}", file=sys.stderr)
        return None


def _md_to_docx(doc, md_text: str, media_dir: Path | None = None):
    """Minimal markdown -> docx: headings, bullets, fenced code, tables, paragraphs, and images."""
    lines = md_text.splitlines()
    in_code = False
    table_lines = []
    
    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        parsed_rows = []
        for line in table_lines:
            parts = [p.strip() for p in line.strip().strip("|").split("|")]
            parsed_rows.append(parts)
            
        if len(parsed_rows) >= 2:
            is_sep = all(re.match(r"^:?-+:?$", cell) for cell in parsed_rows[1]) if parsed_rows[1] else False
            if is_sep:
                headers = parsed_rows[0]
                rows = parsed_rows[2:]
            else:
                headers = parsed_rows[0]
                rows = parsed_rows[1:]
                
            num_cols = len(headers)
            t = doc.add_table(rows=1, cols=num_cols)
            t.style = "Light Grid Accent 1"
            
            for col_idx, h in enumerate(headers):
                if col_idx < len(t.rows[0].cells):
                    t.rows[0].cells[col_idx].text = h
                    
            for r_data in rows:
                row_cells = t.add_row().cells
                for col_idx, val in enumerate(r_data):
                    if col_idx < len(row_cells):
                        row_cells[col_idx].text = val
        table_lines = []

    for line in lines:
        if '<div class="chart-container"' in line or 'data:image/png;base64' in line or 'alt="Figures"' in line or 'alt="Chart"' in line or '</div>' in line:
            continue
        if line.strip().startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(line)
            p.style = doc.styles["No Spacing"]
            for r in p.runs:
                r.font.name = "Courier New"
                r.font.size = Pt(8)
            continue
            
        # Check if line is a table line
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        else:
            if table_lines:
                flush_table()
                
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
            continue
        if re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", line), style="List Bullet")
            continue
        if line.strip():
            # Check if it is a standalone image
            img_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line.strip())
            if img_match:
                alt = img_match.group(1)
                url = img_match.group(2)
                if media_dir:
                    local_path = download_image_cached(url, media_dir)
                    if local_path and local_path.exists():
                        try:
                            doc.add_picture(str(local_path), width=Inches(5.5))
                            continue
                        except Exception as ex:
                            print(f"Failed to insert picture {local_path} into DOCX: {ex}", file=sys.stderr)
                # Fallback to paragraph link
                p = doc.add_paragraph()
                p.add_run(f"📷 [Image: {alt or 'Diagram'}] ").bold = True
                p.add_run(url).italic = True
                continue
                
            doc.add_paragraph(line.strip())
            
    if table_lines:
        flush_table()


# ---- 5c. PDF (reportlab primary, pandoc fallback) -------------------------
def render_pdf(presentable: dict[str, Any], out_path: Path, html_path: Path | None) -> bool:
    fmt_cfg = CONFIG.get("formats", {})
    engine_order = fmt_cfg.get("pdf_engine_order", ["reportlab", "pandoc"])
    
    for engine in engine_order:
        if engine == "reportlab":
            if _HAVE_REPORTLAB and _pdf_via_reportlab(presentable, out_path):
                return True
        elif engine == "pandoc":
            if _HAVE_PANDOC and html_path and html_path.exists():
                rc = os.system(f"pandoc {html_path} -o {out_path} >/dev/null 2>&1")
                if rc == 0 and out_path.exists():
                    return True
    return False


def _pdf_via_reportlab(presentable: dict[str, Any], out_path: Path) -> bool:
    try:
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle("Banner", parent=styles["Normal"],
                                  textColor=_rl_colors.HexColor("#b06a00"),
                                  fontName="Helvetica-Bold", spaceAfter=8))
        flow = []
        flow.append(Paragraph(_html.escape(presentable["title"]), styles["Title"]))
        flow.append(Paragraph(
            f"Session {_html.escape(str(presentable['session_id']))} · "
            f"generated {presentable['generated_at']} · status: {presentable['status']}",
            styles["Italic"]))
        flow.append(Spacer(1, 0.15 * inch))

        banner = presentable.get("banner")
        if banner:
            flow.append(Paragraph(_html.escape(banner), styles["Banner"]))

        # Executive Summary
        exec_sum = presentable.get("executive_summary")
        if exec_sum:
            flow.append(Paragraph("Executive Summary", styles["Heading1"]))
            flow.append(Paragraph(_html.escape(exec_sum), styles["Normal"]))
            flow.append(Spacer(1, 0.12 * inch))

        # Key Takeaways
        takeaways = presentable.get("key_takeaways")
        if takeaways:
            flow.append(Paragraph("Key Takeaways", styles["Heading2"]))
            for item in takeaways:
                flow.append(Paragraph("• " + _html.escape(item), styles["Normal"]))
            flow.append(Spacer(1, 0.12 * inch))

        # Sections
        for sec in presentable.get("sections", []):
            flow.append(Paragraph(_html.escape(sec["title"]), styles["Heading1"]))
            table_lines = []
            
            def flush_pdf_table():
                nonlocal table_lines
                if not table_lines:
                    return
                parsed_rows = []
                for line in table_lines:
                    parts = [p.strip() for p in line.strip().strip("|").split("|")]
                    parsed_rows.append(parts)
                if len(parsed_rows) >= 2:
                    is_sep = all(re.match(r"^:?-+:?$", cell) for cell in parsed_rows[1]) if parsed_rows[1] else False
                    if is_sep:
                        headers = parsed_rows[0]
                        rows = parsed_rows[2:]
                    else:
                        headers = parsed_rows[0]
                        rows = parsed_rows[1:]
                    
                    tdata = [headers]
                    for r in rows:
                        tdata.append(r)
                    flow.append(_rl_table(tdata))
                    flow.append(Spacer(1, 0.12 * inch))
                table_lines = []

            for line in sec.get("body", "").splitlines():
                if '<div class="chart-container"' in line or 'data:image/png;base64' in line or 'alt="Figures"' in line or 'alt="Chart"' in line or '</div>' in line:
                    continue
                if line.strip().startswith("|") and line.strip().endswith("|"):
                    table_lines.append(line)
                    continue
                else:
                    if table_lines:
                        flush_pdf_table()

                m = re.match(r"^(#{1,6})\s+(.*)", line)
                if m:
                    lvl = min(len(m.group(1)), 4)
                    flow.append(Paragraph(_html.escape(m.group(2)),
                                          styles[f"Heading{lvl}"]))
                elif re.match(r"^\s*[-*]\s+", line):
                    flow.append(Paragraph("• " + _html.escape(
                        re.sub(r"^\s*[-*]\s+", "", line)), styles["Normal"]))
                elif line.strip() and not line.strip().startswith("```"):
                    flow.append(Paragraph(_html.escape(line.strip()), styles["Normal"]))
            
            if table_lines:
                flush_pdf_table()
                
            if sec.get("chart_bytes"):
                try:
                    import io
                    img_data = io.BytesIO(sec["chart_bytes"])
                    # Use RLImage imported as RLImage in imports
                    flow.append(RLImage(img_data, width=6*inch, height=2.5*inch))
                    flow.append(Spacer(1, 0.12 * inch))
                except Exception as e:
                    print(f"  [PDF] Failed to add chart image: {e}", file=sys.stderr)
                    
            flow.append(Spacer(1, 0.12 * inch))

        # Tables
        tables = presentable.get("tables", [])
        if tables:
            flow.append(Paragraph("Extracted Parameter Matrix", styles["Heading1"]))
            for table in tables:
                flow.append(Paragraph(_html.escape(table["title"]), styles["Heading2"]))
                headers = table.get("headers", [])
                rows = table.get("rows", [])
                if not rows:
                    continue
                tdata = [headers]
                for row in rows:
                    tdata.append([str(cell)[:40] for cell in row])
                flow.append(_rl_table(tdata))
                flow.append(Spacer(1, 0.12 * inch))

        # Gaps
        gaps = presentable.get("gaps", [])
        if gaps:
            flow.append(Paragraph("Coverage Gaps", styles["Heading1"]))
            for gap in gaps:
                flow.append(Paragraph(f"• <b>{_html.escape(gap['topic'])}</b>: {_html.escape(gap['reason'])}", styles["Normal"]))
            flow.append(Spacer(1, 0.12 * inch))

        # Sources
        sources = presentable.get("sources", [])
        if sources:
            flow.append(Paragraph("Sources & References", styles["Heading1"]))
            tdata = [["Title", "Domain", "URL"]]
            for s in sources:
                tdata.append([
                    str(s.get("title") or "")[:40],
                    str(s.get("domain") or "")[:20],
                    str(s.get("url") or "")[:30]
                ])
            flow.append(_rl_table(tdata))

        try:
            SimpleDocTemplate(str(out_path), pagesize=LETTER,
                              title=presentable["title"]).build(flow)
            return out_path.exists()
        except PermissionError:
            print(f"Error: Permission denied when writing PDF to {out_path}. Please close the file if it is open in another application.")
            fallback_path = out_path.with_name(f"{out_path.stem}_fallback{out_path.suffix}")
            try:
                SimpleDocTemplate(str(fallback_path), pagesize=LETTER,
                                  title=presentable["title"]).build(flow)
                print(f"Successfully saved PDF to fallback path: {fallback_path}")
                return True
            except Exception as ex:
                print(f"Failed to write PDF fallback: {ex}")
                raise
    except Exception as e:
        print(f"  [pdf] reportlab failed: {e}", file=sys.stderr)
        return False


def _rl_table(data):
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _rl_colors.HexColor("#f2f4f7")),
        ("GRID", (0, 0), (-1, -1), 0.4, _rl_colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return t


# ---- 5c. XLSX -------------------------------------------------------------
def render_xlsx(presentable: dict[str, Any], out_path: Path) -> bool:
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("  [xlsx] openpyxl is not installed; skipped.", file=sys.stderr)
        return False

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # Remove default sheet

    # Colors & Fonts
    navy_fill = PatternFill(start_color="1E1B4B", end_color="1E1B4B", fill_type="solid")
    zebra_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    
    white_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    bold_font = Font(name="Calibri", size=11, bold=True)
    normal_font = Font(name="Calibri", size=11)
    
    thin_border = Border(
        left=Side(style='thin', color='E2E8F0'),
        right=Side(style='thin', color='E2E8F0'),
        top=Side(style='thin', color='E2E8F0'),
        bottom=Side(style='thin', color='E2E8F0')
    )

    tables = presentable.get("tables", [])
    
    # 1. Create a sheet for each dynamic table
    for t_idx, table in enumerate(tables):
        title = table.get("title", f"Table {t_idx+1}")
        clean_title = re.sub(r'[\\*?:/\[\]]', '', title)[:30].strip()
        if not clean_title:
            clean_title = f"Table_{t_idx+1}"
        
        orig_title = clean_title
        suffix = 1
        while clean_title in wb.sheetnames:
            clean_title = f"{orig_title[:27]}_{suffix}"
            suffix += 1

        ws = wb.create_sheet(title=clean_title)
        ws.views.sheetView[0].showGridLines = True

        headers = table.get("headers", [])
        rows = table.get("rows", [])

        # Write headers
        for c_idx, h in enumerate(headers):
            cell = ws.cell(row=1, column=c_idx+1, value=h)
            cell.font = white_font
            cell.fill = navy_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border

        # Write rows
        for r_idx, row in enumerate(rows):
            row_num = 2 + r_idx
            fill = zebra_fill if r_idx % 2 == 1 else PatternFill(fill_type=None)
            
            for c_idx, val in enumerate(row):
                col_num = c_idx + 1
                cell = ws.cell(row=row_num, column=col_num, value=val)
                cell.font = normal_font
                cell.border = thin_border
                if fill.fill_type:
                    cell.fill = fill
                    
                if isinstance(val, (int, float)):
                    cell.alignment = Alignment(horizontal="right")
                elif str(val).startswith("http://") or str(val).startswith("https://"):
                    cell.font = Font(name="Calibri", size=11, color="0563C1", underline="single")
                    cell.hyperlink = val
                else:
                    if len(str(val)) > 40:
                        cell.alignment = Alignment(wrap_text=True)

        # Auto-fit columns
        for col in ws.columns:
            max_len = 0
            for cell in col:
                val_str = str(cell.value or '')
                lines = val_str.split('\n')
                for line in lines:
                    max_len = max(max_len, len(line))
            col_letter = get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = min(max(max_len + 3, 10), 60)

    # 2. Sources Sheet
    sources = presentable.get("sources", [])
    if sources:
        ws_src = wb.create_sheet(title="Sources")
        ws_src.views.sheetView[0].showGridLines = True
        headers_src = ["#", "Source Title", "Domain", "Reference URL", "Tier"]
        for c_idx, h in enumerate(headers_src):
            cell = ws_src.cell(row=1, column=c_idx+1, value=h)
            cell.font = white_font
            cell.fill = navy_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border

        for r_idx, src in enumerate(sources):
            row_num = 2 + r_idx
            fill = zebra_fill if r_idx % 2 == 1 else PatternFill(fill_type=None)
            
            ws_src.cell(row=row_num, column=1, value=r_idx+1).font = bold_font
            ws_src.cell(row=row_num, column=2, value=src.get("title") or "").font = normal_font
            ws_src.cell(row=row_num, column=3, value=src.get("domain") or "").font = normal_font
            
            url_cell = ws_src.cell(row=row_num, column=4, value=src.get("url") or "")
            url_cell.font = Font(name="Calibri", size=11, color="0563C1", underline="single")
            if src.get("url"):
                url_cell.hyperlink = src.get("url")
                
            ws_src.cell(row=row_num, column=5, value=src.get("tier") or "").font = normal_font

            for col in range(1, 6):
                cell = ws_src.cell(row=row_num, column=col)
                cell.border = thin_border
                if fill.fill_type:
                    cell.fill = fill

        for col in ws_src.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = get_column_letter(col[0].column)
            ws_src.column_dimensions[col_letter].width = min(max(max_len + 3, 10), 60)

    # 3. Coverage Gaps Sheet
    gaps = presentable.get("gaps", [])
    if gaps:
        ws_gaps = wb.create_sheet(title="Coverage_Gaps")
        ws_gaps.views.sheetView[0].showGridLines = True
        ws_gaps.cell(row=1, column=1, value="Vector / Topic").font = white_font
        ws_gaps.cell(row=1, column=1).fill = navy_fill
        ws_gaps.cell(row=1, column=1).border = thin_border
        ws_gaps.cell(row=1, column=2, value="Reason").font = white_font
        ws_gaps.cell(row=1, column=2).fill = navy_fill
        ws_gaps.cell(row=1, column=2).border = thin_border

        for r_idx, gap in enumerate(gaps):
            row_num = 2 + r_idx
            fill = zebra_fill if r_idx % 2 == 1 else PatternFill(fill_type=None)
            
            ws_gaps.cell(row=row_num, column=1, value=gap.get("topic") or gap.get("vector_id") or "").font = bold_font
            ws_gaps.cell(row=row_num, column=2, value=gap.get("reason") or "Insufficient usable data captured.").font = normal_font

            for col in range(1, 3):
                cell = ws_gaps.cell(row=row_num, column=col)
                cell.border = thin_border
                if fill.fill_type:
                    cell.fill = fill

        ws_gaps.column_dimensions["A"].width = 30
        ws_gaps.column_dimensions["B"].width = 50

    if not wb.sheetnames:
        wb.create_sheet(title="Report")

    try:
        wb.save(str(out_path))
        return True
    except Exception as e:
        print(f"Error saving Excel workbook to {out_path}: {e}", file=sys.stderr)
        return False


# =============================================================================
# 6. CLI
# =============================================================================
def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Deterministic presentation layer "
                                             "(Stage 3, no API).")
    ap.add_argument("run_dir", help="Path to run_<id> folder")
    ap.add_argument("--formats", default=None,
                    help="comma list: docx,html,pdf (default: from run_config "
                         "output_format, else all)")
    ap.add_argument("--no-charts", action="store_true",
                    help="render tables instead of PNG charts")
    args = ap.parse_args(argv)

    run_dir = Path(args.run_dir).expanduser().resolve()
    if not run_dir.is_dir():
        print(f"ERROR: not a directory: {run_dir}", file=sys.stderr)
        return 2

    # Load best payload
    payload = load_best_payload(run_dir)
    
    # Load sources
    sources = load_json_if_exists(run_dir / "sources.json", [])
    
    # Load config
    config = load_present_config()
    
    # Run through presentation filter to build presentable allowlist payload
    import presentation_filter
    presentable = presentation_filter.build_presentable(payload, sources, config)

    # Matplotlib chart rendering and embedding
    if not args.no_charts and _HAVE_MPL:
        # Match presentable sections with payload sections to access data
        payload_sections_by_id = {sec.get("vector_id"): sec for sec in payload.get("sections", []) if sec.get("vector_id")}
        
        for p_sec in presentable.get("sections", []):
            vid = p_sec.get("vector_id")
            if not vid:
                continue
            matching_payload_sec = payload_sections_by_id.get(vid)
            if not matching_payload_sec:
                continue
            
            sec_data = matching_payload_sec.get("data")
            if not sec_data:
                continue
                
            # Detect visuals (like bar charts) in the extracted section data
            visuals = detect_visuals(sec_data)
            for vis in visuals:
                if vis.get("kind") == "barchart":
                    png_bytes = render_barchart(vis)
                    if png_bytes:
                        b64_str = base64.b64encode(png_bytes).decode("utf-8")
                        img_tag = f'\n\n<div class="chart-container" style="text-align: center; margin: 1.5rem 0;"><img src="data:image/png;base64,{b64_str}" alt="{vis.get("title", "Chart")}" style="max-width: 100%; height: auto; border: 1px solid var(--border); border-radius: var(--radius-md);"/></div>\n\n'
                        # Append to markdown body
                        p_sec["body"] = p_sec.get("body", "") + img_tag
                        p_sec["chart_bytes"] = png_bytes
                        print(f"  [Charts] Embedded Matplotlib bar chart for vector '{vid}' into section body.")

    # Stamp banner if fallback mode is active
    banner_cfg = config.get("banner", {})
    fallback_text = banner_cfg.get("fallback_text", "⚠️ MECHANICALLY ASSEMBLED — no AI synthesis ran; structured from extracted data only.")
    trigger_modes = banner_cfg.get("trigger_synthesis_modes", ["fallback_no_api", "incomplete_fallback"])
    
    if payload.get("is_fallback") or payload.get("synthesis_mode") in trigger_modes or presentable.get("status") == "incomplete_fallback":
        presentable["banner"] = fallback_text

    # Render HTML template
    html_out = render_html_template(presentable)

    # Quality Gate Check
    assert_no_template_or_placeholder_leak(html_out)

    # decide formats
    if args.formats:
        formats = [f.strip().lower() for f in args.formats.split(",") if f.strip()]
    else:
        run_config = load_json_if_exists(run_dir / "run_config.json", {})
        of = (run_config.get("output_format") or "").lower()
        if of == "excel":
            of = "xlsx"
        fmt_cfg = CONFIG.get("formats", {})
        default_formats = fmt_cfg.get("default", ["docx", "html", "pdf"])
        formats = [of] if of in ("docx", "html", "pdf", "xlsx") else default_formats
        if "html" not in formats and "xlsx" not in formats:
            formats.append("html")

    written = []
    html_path = run_dir / "report.html"
    if "html" in formats:
        try:
            html_path.write_text(html_out, encoding="utf-8")
            written.append(html_path.name)
        except PermissionError:
            print(f"Error: Permission denied when writing to {html_path}.")
            fallback_html = html_path.with_name("report_fallback.html")
            fallback_html.write_text(html_out, encoding="utf-8")
            written.append(fallback_html.name)
            
    if "docx" in formats:
        p = run_dir / "report.docx"
        if render_docx(presentable, p):
            written.append(p.name)
        else:
            print("  [docx] python-docx unavailable; skipped.", file=sys.stderr)
            
    if "xlsx" in formats:
        p = run_dir / "report.xlsx"
        if render_xlsx(presentable, p):
            written.append(p.name)
        else:
            print("  [xlsx] openpyxl unavailable; skipped.", file=sys.stderr)
            
    if "pdf" in formats:
        if not html_path.exists():
            try:
                html_path.write_text(html_out, encoding="utf-8")
            except PermissionError:
                pass
        p = run_dir / "report.pdf"
        if render_pdf(presentable, p, html_path):
            written.append(p.name)
        else:
            print("  [pdf] no working PDF engine; skipped (html/docx still produced).",
                  file=sys.stderr)

    print(f"Presented {len(presentable['sections'])} sections and {len(presentable['tables'])} tables.")
    print("Wrote: " + ", ".join(written) if written else "Wrote: (nothing)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())